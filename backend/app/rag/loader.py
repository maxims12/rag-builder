"""Local file loader: walk ``docs/`` and parse files into ParsedDocuments.

This is the local-source implementation of the ``DocumentLoader`` interface from
CONTRACT.md §4. It walks the configured folder (honouring ``SourcesConfig``),
parses each supported file type into clean text, and emits ``ParsedDocument``
records whose ``metadata`` carries the contract-mandated local-file fields:

    source, file_name, file_type, modified_at (ISO 8601 UTC), content_hash

The web loader (Phase 5) emits the *same* ``ParsedDocument`` shape (with the
web metadata variant) so everything downstream — chunk → embed → store — is
shared and reused unchanged.
"""

from __future__ import annotations

import csv
import fnmatch
import hashlib
import io
import logging
import os
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, List

from pydantic import BaseModel, Field

from app.config_schemas import SourcesConfig

logger = logging.getLogger("app.rag.loader")


class ParsedDocument(BaseModel):
    """One parsed source document (file or web page) with its metadata.

    For local files, ``metadata`` includes:
      - ``source``: str (file path)
      - ``file_name``: str (base name)
      - ``file_type``: str (extension incl. dot, e.g. '.pdf')
      - ``modified_at``: str (ISO 8601 UTC mtime)
      - ``content_hash``: str (SHA-256 of text)

    For web sources (Phase 5), ``metadata`` includes:
      - ``source_url``, ``title``, ``fetched_at``, ``content_hash``
    """

    text: str
    metadata: Dict[str, Any] = Field(default_factory=dict)


def compute_content_hash(text: str) -> str:
    """SHA-256 hex digest of the (UTF-8 encoded) text contents."""
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


def _iso_utc(ts: float) -> str:
    return datetime.fromtimestamp(ts, tz=timezone.utc).isoformat()


# ── Per-extension parsers ─────────────────────────────────────────────
# Each returns clean plain text. Heavy/optional imports are done lazily so the
# module imports even if a parser's backing library is unavailable.


def _parse_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def _parse_md(path: Path) -> str:
    # Markdown is already human-readable text; keep the raw source so headings
    # survive for the markdown chunking strategy.
    return path.read_text(encoding="utf-8", errors="replace")


def _parse_pdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    parts: List[str] = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception as exc:  # pragma: no cover - corrupt page
            logger.warning("PDF page extraction failed for %s: %s", path, exc)
    return "\n\n".join(p for p in parts if p.strip())


def _parse_docx(path: Path) -> str:
    import docx  # python-docx

    document = docx.Document(str(path))
    return "\n".join(p.text for p in document.paragraphs if p.text.strip())


def _parse_html(path: Path) -> str:
    from bs4 import BeautifulSoup

    raw = path.read_text(encoding="utf-8", errors="replace")
    soup = BeautifulSoup(raw, "html.parser")
    # Drop boilerplate that pollutes RAG context.
    for tag in soup(["script", "style", "nav", "header", "footer", "noscript"]):
        tag.decompose()
    return soup.get_text(separator="\n", strip=True)


def _parse_csv(path: Path) -> str:
    raw = path.read_text(encoding="utf-8", errors="replace")
    out = io.StringIO()
    reader = csv.reader(io.StringIO(raw))
    for row in reader:
        out.write(", ".join(cell.strip() for cell in row))
        out.write("\n")
    return out.getvalue()


# Extension (lower-case, incl. dot) -> parser function.
PARSERS: Dict[str, Any] = {
    ".txt": _parse_txt,
    ".md": _parse_md,
    ".markdown": _parse_md,
    ".pdf": _parse_pdf,
    ".docx": _parse_docx,
    ".html": _parse_html,
    ".htm": _parse_html,
    ".csv": _parse_csv,
}


def parse_file(path: Path) -> str:
    """Parse a single file into clean text. Raises if the type is unsupported."""
    ext = path.suffix.lower()
    parser = PARSERS.get(ext)
    if parser is None:
        raise ValueError(f"Unsupported file type: {ext}")
    return parser(path)


class LocalFileLoader:
    """``DocumentLoader`` for the local ``docs/`` folder.

    Honours ``SourcesConfig``: ``docs_folder``, ``recursive``, ``file_types``,
    ``exclude_patterns`` (glob), and ``max_file_size_mb``.
    """

    def __init__(self, config: SourcesConfig) -> None:
        self.config = config

    def _allowed_extensions(self) -> set[str]:
        return {ext.lower() for ext in self.config.file_types}

    def _is_excluded(self, rel_path: str) -> bool:
        for pattern in self.config.exclude_patterns:
            if fnmatch.fnmatch(rel_path, pattern) or fnmatch.fnmatch(
                os.path.basename(rel_path), pattern
            ):
                return True
        return False

    def discover(self) -> List[Path]:
        """Return the list of candidate file paths to parse."""
        base = Path(self.config.docs_folder)
        if not base.exists():
            logger.warning("docs folder does not exist: %s", base)
            return []

        allowed = self._allowed_extensions()
        max_bytes = self.config.max_file_size_mb * 1024 * 1024
        files: List[Path] = []

        iterator = base.rglob("*") if self.config.recursive else base.glob("*")
        for path in iterator:
            if not path.is_file():
                continue
            if path.suffix.lower() not in allowed:
                continue
            try:
                rel = str(path.relative_to(base))
            except ValueError:
                rel = path.name
            if self._is_excluded(rel):
                continue
            try:
                if path.stat().st_size > max_bytes:
                    logger.warning(
                        "Skipping %s: exceeds max_file_size_mb=%s",
                        path,
                        self.config.max_file_size_mb,
                    )
                    continue
            except OSError:
                continue
            files.append(path)
        return sorted(files)

    def load_path(self, path: Path) -> ParsedDocument | None:
        """Parse a single discovered file into a ParsedDocument (or None)."""
        try:
            text = parse_file(path)
        except Exception as exc:
            logger.warning("Failed to parse %s: %s", path, exc)
            return None
        if not text.strip():
            logger.info("Skipping empty file: %s", path)
            return None
        try:
            mtime = path.stat().st_mtime
        except OSError:
            mtime = datetime.now(timezone.utc).timestamp()
        return ParsedDocument(
            text=text,
            metadata={
                "source": str(path),
                "file_name": path.name,
                "file_type": path.suffix.lower(),
                "modified_at": _iso_utc(mtime),
                "content_hash": compute_content_hash(text),
            },
        )

    async def load(self) -> List[ParsedDocument]:
        """Discover, parse, and return all documents from the configured folder."""
        docs: List[ParsedDocument] = []
        for path in self.discover():
            parsed = self.load_path(path)
            if parsed is not None:
                docs.append(parsed)
        logger.info("LocalFileLoader loaded %d document(s)", len(docs))
        return docs
